# Librerías
import tkinter as tk
from tkinter import messagebox
from PIL import Image, ImageTk
import os
import math
import numpy as np
import pandas as pd
from docx import Document
import PyPDF2
from itertools import product
from itertools import combinations
from tkinter import simpledialog
import copy
import matplotlib.pyplot as plt
from tkinter import filedialog

# Funciones de ejemplo para las opciones
def volver_a_principal():
    """
    Clears the frame, updates the content label, and packs the top bar and About button.
    """
    limpiar_frame()
    contenido_label.config(text="Análisis de probabilidades")
    barra_superior.pack(side=tk.TOP, fill=tk.X)

    About_Button.pack(side=tk.LEFT)
# Funcion para la opcion 1
def opcion1():
    """
    Function to handle option 1.
    Clears the frame, updates the label text, and adds buttons for matrix input and loading.
    """
    limpiar_frame()
    contenido_label.config(text="Ingreso de matrices")
    barra_superior.pack(side=tk.TOP, fill=tk.X)

    Ingresar_Button = tk.Button(barra_superior, text="Ingresar Matriz", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=Ingresar_B)
    Ingresar_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    Cargar_Button = tk.Button(barra_superior, text="Cargar matriz", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10), command=Cargar_B)
    Cargar_Button.pack(side=tk.LEFT)  # Alinear a la izquierda
# Funcion para la opcion 2
def opcion2():
    """
    Function to handle option 2.
    Clears the frame, updates the label, and adds buttons for matrix analysis.
    """
    limpiar_frame()
    contenido_label.config(text="Análisis de matrices asociadas")
    barra_superior.pack(side=tk.TOP, fill=tk.X)

    Matriz_Button = tk.Button(barra_superior, text="Matriz Original", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10), command=MatrizO_B)
    Matriz_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    Matriz_CanalE_Button = tk.Button(barra_superior, text="Matriz Canal-Estado", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=MatrizC_B)
    Matriz_CanalE_Button.pack(side=tk.LEFT)  # Alinear a la izquierda    

    Matriz_EstadoE_Button = tk.Button(barra_superior, text="Matriz Estado-Estado", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=MatrizE_B)
    Matriz_EstadoE_Button.pack(side=tk.LEFT)  # Alinear a la izquierda
# Funcion para la opcion 3
def opcion3():
    """
    Function to handle option 3.
    Clears the frame, updates the label text, and adds buttons for entering matrix state, marginalizing, and viewing a graph.
    """
    limpiar_frame()
    contenido_label.config(text="Visualización y probabilidades")
    barra_superior.pack(side=tk.TOP, fill=tk.X)

    Ingresar_MatrizE_Button = tk.Button(barra_superior, text="Ingresar Matriz Estado", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=IngMatrizE_B)
    Ingresar_MatrizE_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    Marginalizar_Button = tk.Button(barra_superior, text="Marginalizar", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=Marginalizar_B)
    Marginalizar_Button.pack(side=tk.LEFT)  # Alinear a la izquierda    

    Grafica_Button = tk.Button(barra_superior, text="Ver Gráfica", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=Grafica_B)
    Grafica_Button.pack(side=tk.LEFT)  # Alinear a la izquierda
# Funcion para la opcion 4
def opcion4():
    """
    Function to handle Option 4.

    This function clears the frame, updates the content label, and adds buttons for various actions.

    Args:
        None

    Returns:
        None
    """
    limpiar_frame()
    contenido_label.config(text="Distancia EMD")
    barra_superior.pack(side=tk.TOP, fill=tk.X)

    Ingresar_MatrizEC_Button = tk.Button(barra_superior, text="Ingresar Matriz Estado", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=IngMatrizE_B)
    Ingresar_MatrizEC_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    DistanciaEMD_Button = tk.Button(barra_superior, text="Distancia EMD", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=DistanciaE_B)
    DistanciaEMD_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    GraficaO_Button = tk.Button(barra_superior, text="Gráfica Original", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=GraficaO_B)
    GraficaO_Button.pack(side=tk.LEFT)  # Alinear a la izquierda

    GraficaG_Button = tk.Button(barra_superior, text="Gráfica Ganadora", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 10),command=GraficaG_B)
    GraficaG_Button.pack(side=tk.LEFT)  # Alinear a la izquierda
# Funcion para la opcion 5
def salir_del_programa():
    """
    Closes the main window of the program.
    """
    ventana_principal.destroy()
# Funcion que limpia el frame
def limpiar_frame():
    """
    Clears the contents of the frame_trabajo and barra_superior widgets.
    """
    for widget in frame_trabajo.winfo_children():
        widget.pack_forget()
        widget.place_forget()
        widget.grid_forget()

    for widget in barra_superior.winfo_children():
        widget.pack_forget()
        widget.place_forget()
        widget.grid_forget()
# Funcion que limpia el frame tipo Tr
def limpiar_frameTr():
    """
    Clears all the widgets inside the frame_trabajo.
    """
    for widget in frame_trabajo.winfo_children():
        widget.pack_forget()
        widget.place_forget()
        widget.grid_forget()
# Lista para canales
channels=[]
# Función que implementa el about
def aboutB():
    """
    Display information about the project and its members.
    """
    messagebox.showinfo("Acerca de", "Proyecto - Análisis\n\nIntegrantes:\n\n- Juan Esteban Graell Alzate\n- Juan Camilo Toro Palacio\n\n Presentado a: Luz Enith Guerrero Mendietai")
# Función que imprime las tablas
def imprimir_tabla(matriz):
    """
    Imprime una tabla en la interfaz gráfica con los datos de una matriz.

    Args:
        matriz (list): La matriz de datos a mostrar en la tabla.

    Raises:
        ValueError: Si la matriz está vacía.

    Returns:
        None
    """
    try:
        if len(matriz) > 0:  # Asegurarse de que la matriz no esté vacía
            limpiar_frameTr()
            filas = len(matriz) + 1
            columnas = len(matriz[0]) + 1

            # Crear etiqueta vacía para la celda en la fila 0 y la columna 0
            tk.Label(frame_trabajo, text="",bg="#353535", font=('Helvetica', 15)).grid(row=0, column=0, sticky="nsew", padx=1, pady=1)

            # Crear etiquetas para las columnas
            for j in range(1, columnas):
                tk.Label(frame_trabajo, text="T"+str(j), bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=0, column=j, sticky="nsew", padx=1, pady=1)

            # Crear etiquetas para las filas
            for i in range(1, filas):
                tk.Label(frame_trabajo, text=diccionario_letras[i],  bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=i, column=0, sticky="nsew", padx=1, pady=1)

            # Rellenar la tabla con los datos de la matriz
            for i in range(1, filas):
                for j in range(1, columnas):
                    tk.Label(frame_trabajo, text=str(matriz[i-1][j-1]), bg="#6F6F6F", fg="white", font=('Helvetica', 15)).grid(row=i, column=j, sticky="nsew", padx=1, pady=1)

            # Configurar el peso de las filas y columnas para hacer la tabla expansible
            for i in range(filas):
                frame_trabajo.grid_rowconfigure(i)

            for j in range(columnas):
                frame_trabajo.grid_columnconfigure(j)
        else:
            raise ValueError("La matriz está vacía.")
    except Exception as e:
        root = tk.Tk()
        root.withdraw()  # Para evitar que se muestre la ventana Tk vacía
        messagebox.showerror("Error", str(e))
        root.destroy()
# Función que imprime las tablas tipo CS
def imprimir_tablaCS(matriz):
    """
    Imprime una tabla en la interfaz gráfica con los datos de una matriz.

    Args:
        matriz (list): La matriz de datos a imprimir en la tabla.

    Raises:
        ValueError: Si la matriz está vacía.

    """
    try:
        if len(matriz) > 0:  # Asegurarse de que la matriz no esté vacía
            limpiar_frameTr()
            filas = len(matriz)
            columnas = len(matriz[0])

            # Crear etiquetas para las columnas
            for j in range(columnas):
                tk.Label(frame_trabajo, text=matriz[0][j], bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=0, column=j, sticky="nsew", padx=1, pady=1)

            # Crear etiquetas para las filas y rellenar la tabla con los datos de la matriz
            for i in range(1, filas):
                tk.Label(frame_trabajo, text=matriz[i][0],  bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=i, column=0, sticky="nsew", padx=1, pady=1)
                for j in range(1, min(columnas, len(matriz[i]))):
                    if matriz[i][j] is not None:
                        tk.Label(frame_trabajo, text=str(matriz[i][j]), bg="#6F6F6F", fg="white", font=('Helvetica', 15)).grid(row=i, column=j, sticky="nsew", padx=1, pady=1)

            # Configurar el peso de las filas y columnas para hacer la tabla expansible
            for i in range(filas):
                frame_trabajo.grid_rowconfigure(i)

            for j in range(columnas):
                frame_trabajo.grid_columnconfigure(j)
        else:
            raise ValueError("La matriz está vacía.")
    except Exception as e:
        root = tk.Tk()
        root.withdraw()  # Para evitar que se muestre la ventana Tk vacía
        messagebox.showerror("Error", str(e))
        root.destroy()
# Función que imprime las tablas tipo S
def imprimir_tablaS(matriz):
    """
    Imprime una tabla en la interfaz gráfica con los datos de una matriz.

    Args:
        matriz (list): La matriz de datos a imprimir en la tabla.

    Raises:
        ValueError: Si la matriz está vacía.

    """
    try:
        if len(matriz) > 0:  # Asegurarse de que la matriz no esté vacía
            limpiar_frameTr()
            filas = len(matriz)
            columnas = len(matriz[0])

            # Crear etiquetas para las columnas
            for j in range(columnas):
                tk.Label(frame_trabajo, text=matriz[0][j], bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=0, column=j, sticky="nsew", padx=1, pady=1)

            # Crear etiquetas para las filas y rellenar la tabla con los datos de la matriz
            for i in range(1, filas):
                tk.Label(frame_trabajo, text=matriz[i][0],  bg="#353535", fg="white", font=('Helvetica', 15)).grid(row=i, column=0, sticky="nsew", padx=1, pady=1)
                for j in range(1, min(columnas-1, len(matriz[i]))+1):
                    if matriz[i][j] is not None:
                        tk.Label(frame_trabajo, text=str(matriz[i][j]), bg="#6F6F6F", fg="white", font=('Helvetica', 15)).grid(row=i, column=j, sticky="nsew", padx=1, pady=1)

            # Configurar el peso de las filas y columnas para hacer la tabla expansible
            for i in range(filas):
                frame_trabajo.grid_rowconfigure(i)

            for j in range(columnas-1):
                frame_trabajo.grid_columnconfigure(j)
        else:
            raise ValueError("La matriz está vacía.")
    except Exception as e:
        root = tk.Tk()
        root.withdraw()  # Para evitar que se muestre la ventana Tk vacía
        messagebox.showerror("Error", str(e))
        root.destroy()
# Función que carga los archivos
def loadFromFile(file_path):
    """
    Load data from a file.

    Args:
        file_path (str): The path to the file.

    Returns:
        list: The loaded data as a matrix.

    Raises:
        FileNotFoundError: If the file is not found.

    """
    _, file_extension = os.path.splitext(file_path)
    try:
        if file_extension == '.txt':
            with open(file_path, 'r') as file:
                matrix = [[int(value) for value in line.strip().split(',')] for line in file]
                return matrix
        elif file_extension == '.xlsx':
            excel_data = pd.read_excel(file_path, header=None)
            return excel_data.values.tolist()
        elif file_extension == '.docx':
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            lines = text.split("\n")
            matrix = [[int(number) for number in line.strip().split() if number.isdigit()] for line in lines if line.strip().split()]
            return matrix
        elif file_extension == '.pdf':
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                lines = [page.extract_text().split('\n') for page in pdf_reader.pages]
                matrix = [list(map(int, line.split())) for sublist in lines for line in sublist if line]
                return matrix
        else:
            messagebox.showerror("Error", "Tipo de archivo no soportado.")
            return None
    except FileNotFoundError:
        messagebox.showerror("Error", f"El archivo {file_path} no se encontró. Asegúrate de que el archivo exista y vuelva a intentarlo.")
        return None
# Función que ayuda a cargar los archivos desde el explorador
def abrir_explorador_y_cargar_archivo():
    """
    Opens a file dialog to select a file and loads its contents into a matrix.

    Returns:
        matrix (list): A matrix containing the contents of the selected file.

    Raises:
        FileNotFoundError: If the selected file is not found.
    """
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), 
                                                      ("Excel files", "*.xlsx"), 
                                                      ("Word files", "*.docx"), 
                                                      ("PDF files", "*.pdf")])
    if file_path:
        _, file_extension = os.path.splitext(file_path)
        try:
            if file_extension == '.txt':
                with open(file_path, 'r') as file:
                    matrix = [[int(value) for value in line.strip().split(',')] for line in file]
                    return matrix
            elif file_extension == '.xlsx':
                excel_data = pd.read_excel(file_path, header=None)
                return excel_data.values.tolist()
            elif file_extension == '.docx':
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                lines = text.split("\n")
                matrix = [[int(number) for number in line.strip().split() if number.isdigit()] for line in lines if line.strip().split()]
                return matrix
            elif file_extension == '.pdf':
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    lines = [page.extract_text().split('\n') for page in pdf_reader.pages]
                    matrix = [list(map(int, line.split())) for sublist in lines for line in sublist if line]
                    return matrix
            else:
                messagebox.showerror("Error", "Tipo de archivo no soportado.")
                return None
        except FileNotFoundError:
            messagebox.showerror("Error", f"El archivo {file_path} no se encontró. Asegúrate de que el archivo exista y vuelva a intentarlo.")
            return None
    else:
        messagebox.showerror("Error", "No se seleccionó ningún archivo.")
        return None
# Función que carga los archivos solo para la matriz estado
def loadFromFileMatrixE(file_path):
    """
    Load a matrix from a file.

    Args:
        file_path (str): The path to the file.

    Returns:
        list: The loaded matrix.

    Raises:
        FileNotFoundError: If the file is not found.

    """
    _, file_extension = os.path.splitext(file_path)
    try:
        if file_extension == '.txt':
            with open(file_path, 'r') as file:
                matrix = [[int(value) for value in line.strip().split(',')] for line in file]
                return matrix
        elif file_extension == '.xlsx':
            excel_data = pd.read_excel(file_path, header=None).values.tolist()
            new_data = []
            for row in excel_data:
                new_row = []
                for item in row:
                    if isinstance(item, str):
                        new_row.append([int(digit) for digit in item.split(',')])
                    else:
                        new_row.append([item])
                new_data.append(new_row)     
            new_data[0].pop(0)

            return new_data
        elif file_extension == '.docx':
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            lines = text.split("\n")
            matrix = [[int(number) for number in line.strip().split() if number.isdigit()] for line in lines if line.strip().split()]
            return matrix
        elif file_extension == '.pdf':
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                lines = [page.extract_text().split('\n') for page in pdf_reader.pages]
                matrix = [list(map(int, line.split())) for sublist in lines for line in sublist if line]
                return matrix
        else:
            messagebox.showerror("Error", "Tipo de archivo no soportado.")
            return None
    except FileNotFoundError:
        messagebox.showerror("Error", f"El archivo {file_path} no se encontró. Asegúrate de que el archivo exista y vuelva a intentarlo.")
        return None
# Función que ayuda a cargar los archivos desde el explorador solo para la matriz estado
def abrir_explorador_y_cargar_archivoMatrixE():
    """
    Opens a file dialog to select a file and loads the contents of the file into a matrix.

    Returns:
        matrix (list): The matrix containing the loaded data.

    Raises:
        FileNotFoundError: If the selected file is not found.
    """
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        _, file_extension = os.path.splitext(file_path)
        try:
            if file_extension == '.txt':
                with open(file_path, 'r') as file:
                    matrix = [[int(value) for value in line.strip().split(',')] for line in file]
                    return matrix
            elif file_extension == '.xlsx':
                excel_data = pd.read_excel(file_path, header=None).values.tolist()
                new_data = []
                for row in excel_data:
                    new_row = []
                    for item in row:
                        if isinstance(item, str):
                            new_row.append([int(digit) for digit in item.split(',')])
                        else:
                            new_row.append([item])
                    new_data.append(new_row)     
                new_data[0].pop(0)

                return new_data
            
            elif file_extension == '.docx':
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                lines = text.split("\n")
                matrix = [[int(number) for number in line.strip().split() if number.isdigit()] for line in lines if line.strip().split()]
                return matrix
            elif file_extension == '.pdf':
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    lines = [page.extract_text().split('\n') for page in pdf_reader.pages]
                    matrix = [list(map(int, line.split())) for sublist in lines for line in sublist if line]
                    return matrix
            else:
                messagebox.showerror("Error", "Tipo de archivo no soportado.")
                return None
        except FileNotFoundError:
            messagebox.showerror("Error", f"El archivo {file_path} no se encontró. Asegúrate de que el archivo exista y vuelva a intentarlo.")
            return None
    else:
        messagebox.showerror("Error", "No se seleccionó ningún archivo.")
        return None
# Función que crea los canales
def create_channels():
    """
    Creates channels based on the data entered in a table.

    Returns:
    - datos_tabla (list): A list containing the data entered in the table.

    Raises:
    - ValueError: If the input in the table is empty or not an integer.
    - messagebox.showerror: If the input in the table is negative.

    """
    global channels
    datos_tabla = []
    for i in range(filasM):  # Comenzar desde 1 para evitar los títulos de las filas
        fila = []
        for j in range(columnasM):  # Comenzar desde 1 para evitar los títulos de las columnas
            try:
                valor = int(tablaM[i][j].get())
                if valor < 0:  # Verificar si el valor es negativo
                    messagebox.showerror("Error", f"La entrada en la fila {i}, columna {j} no puede ser un número negativo.")
                    return
                fila.append(valor)
            except ValueError:
                messagebox.showerror("Error", f"La entrada en la fila {i}, columna {j} está vacía o no es un número entero.")
                return
        datos_tabla.append(fila)
        channels=datos_tabla
    messagebox.showinfo("Éxito", "Matriz insertada exitosamente.")
    volver_a_principal()
    return datos_tabla
# Función que crea tablas
def crear_tabla(filas,columnas):
    """
    Crea una tabla en la interfaz gráfica con el número de filas y columnas especificadas.

    Args:
        filas (int): El número de filas de la tabla.
        columnas (int): El número de columnas de la tabla.

    Returns:
        None
    """
    global filasM, columnasM, frame_trabajo, tablaM
    filasM = int(filas)
    columnasM = int(columnas)

    # Crear títulos para las columnas
    for j in range(columnasM + 1):
        if j == 0:
            label = tk.Label(frame_trabajo, text="    ",bg="#353535", fg="white", font=('Helvetica', 15))
        else:
            label = tk.Label(frame_trabajo, text=f"T{j}",bg="#353535", fg="white", font=('Helvetica', 15))

        label.grid(row=0, column=j, sticky="nsew", padx=1, pady=1)  # sticky para expandir al norte, sur, este y oeste

    # Crear títulos para las filas y las entradas de la tabla
    tablaM = [[] for _ in range(filasM)]
    for i in range(filasM):
        # Título de la fila
        label = tk.Label(frame_trabajo, text=diccionario_letras[i+1],bg="#353535", fg="white", font=('Helvetica', 15))
        label.grid(row=i + 1, column=0, sticky="nsew", padx=1, pady=1)

        for j in range(columnasM):
            # Entradas en la tabla
            entry = tk.Entry(frame_trabajo, width=1,bg="#6F6F6F", fg="white",font=('Helvetica', 15))
            entry.grid(row=i + 1, column=j + 1, sticky="nsew", padx=1, pady=1)
            tablaM[i].append(entry)
        frame_trabajo.grid_rowconfigure(j)

    crearM_Button=tk.Button(frame_trabajo, text="Ingresar", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=create_channels)
    crearM_Button.place(relx=0.5,rely=0.9, anchor="center")
# Función que crea tablas tipo M
def crear_tablaM(filas,columnas):
    """
    Crea una tabla en la interfaz gráfica con el número de filas y columnas especificadas.

    Args:
        filas (int): El número de filas de la tabla.
        columnas (int): El número de columnas de la tabla.

    Returns:
        None
    """
    global filasM, columnasM, frame_trabajo, tablaM
    filasM = int(filas)
    columnasM = int(columnas)

    # Crear títulos para las columnas
    for j in range(columnasM + 1):
        if j == 0:
            label = tk.Label(frame_trabajo, text="    ",bg="#353535", fg="white", font=('Helvetica', 15))
        else:
            label = tk.Label(frame_trabajo, text=f"T{j}",bg="#353535", fg="white", font=('Helvetica', 15))

        label.grid(row=0, column=j, sticky="nsew", padx=1, pady=1)  # sticky para expandir al norte, sur, este y oeste

    # Crear títulos para las filas y las entradas de la tabla
    tablaM = [[] for _ in range(filasM)]
    for i in range(filasM):
        # Título de la fila
        label = tk.Label(frame_trabajo, text=diccionario_letras[i+1],bg="#353535", fg="white", font=('Helvetica', 15))
        label.grid(row=i + 1, column=0, sticky="nsew", padx=1, pady=1)

        for j in range(columnasM):
            # Entradas en la tabla
            entry = tk.Entry(frame_trabajo, width=1,bg="#6F6F6F", fg="white",font=('Helvetica', 15))
            entry.grid(row=i + 1, column=j + 1, sticky="nsew", padx=1, pady=1)
            tablaM[i].append(entry)
        frame_trabajo.grid_rowconfigure(j)            
        

    # Configurar el peso de las columnas para que se expandan uniformemente horizontalmente
    for j in range(columnasM + 1):
        frame_trabajo.grid_columnconfigure(j)

    crearM_Button=tk.Button(frame_trabajo, text="Ingresar", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=create_channels)
    crearM_Button.place(relx=0.5,rely=1, anchor="center")
# Función para invertir una matriz
def invertMatrix(matrix):
    """
    Inverts the given matrix by transposing its rows and columns.

    Args:
        matrix (list): The matrix to be inverted.

    Returns:
        list: The inverted matrix.
    """
    invertedMatrix = []
    for i in range(len(matrix[0])):
        temp = []
        for j in range(len(matrix)):
            temp.append(matrix[j][i])
        invertedMatrix.append(temp)
    return invertedMatrix
# Función para crear una matriz de estados
def createStateMatrix(matrix):
    """
    Create a state matrix based on the given matrix.

    Args:
        matrix (list): The input matrix.

    Returns:
        list: The state matrix.
    """
    stateMatrix = []
    times = 0

    for i in range(len(matrix) - 1, -1, -1):
        state = []
        state[len(state):] = [0] * (2 ** i)
        state[len(state):] = [1] * (2 ** i)
        temp = state
        for i in range(times):
            state[len(state):] = temp
        stateMatrix.append(state)
        times += 1
    return stateMatrix
# Función para crear una matriz de canales
def createChannelMatrix(matrix):
    """
    Creates a channel matrix based on the given matrix.

    Args:
        matrix (list): The input matrix.

    Returns:
        list: The created channel matrix.
    """
    createdMatrix = [["Canales:"]]

    for i in range(len(matrix[0])):
        createdMatrix[0].append("Canal " + str(i + 1))

    for i in range(len(matrix)):
        createdMatrix.append([matrix[i], ])

    return createdMatrix
# Función para simplificar una fracción a formato decimal
def simplify_fraction(numerator, denominator):
    """
    Simplifies a fraction by dividing the numerator and denominator by their greatest common divisor.

    Args:
        numerator (int): The numerator of the fraction.
        denominator (int): The denominator of the fraction.

    Returns:
        float or str: The simplified fraction as a decimal number with two decimal places, or "Undefined" if the denominator is zero.
    """
    if denominator != 0:
        result = numerator / denominator
        return math.trunc(result * 100) / 100  # Truncate to two decimal places
    else:
        return "Undefined"
# Función para crear  la matriz canal estado
def channelState(matrix):
    """
    Calculates the state channel matrix based on the given matrix.

    Args:
        matrix (list): The input matrix.

    Returns:
        list: The state channel matrix.

    Raises:
        ValueError: If the matrix is empty.
    """
    try:
        if len(matrix) > 0:  # Asegurarse de que la matriz no esté vacía
            root = tk.Tk()
            root.withdraw()

            invertedMatrix = invertMatrix(matrix)
            invertedStateMatrix = invertMatrix(createStateMatrix(matrix))
            channels = []

            for i in range(len(matrix)):
                n = simpledialog.askinteger(title="Canal", prompt=f"¿Qué dato desea en el canal {i + 1} para la revisión?")
                while n != 0 and n != 1:
                    tk.messagebox.showerror("Error", "El estado solo acepta números binarios")
                    n = simpledialog.askinteger(title="Canal", prompt=f"¿Qué dato desea en el canal {i + 1} para la revisión?")
                channels.append(n)
            t = simpledialog.askinteger(title="Estado", prompt="¿En qué estado siguiente t+n o previo t-n desea comparar los canales? (Ponga negativo para previos)")

            stateChannelMatrix = createChannelMatrix(invertedStateMatrix)

            for h in range(len(invertedMatrix[0])):
                for i in range(len(invertedStateMatrix)):
                    channelCases = []
                    total, cases = 0, 0
                    for j in range(len(invertedMatrix)-1):
                        if invertedStateMatrix[i] == invertedMatrix[j]:
                            total += 1
                            if j + t >= len(invertedMatrix):
                                if invertedMatrix[t-1][h] == channels[h]:
                                    cases += 1
                            elif invertedMatrix[j + t][h] == channels[h]:
                                cases += 1

                    # Evitar la división por 0
                    if total == 0:
                        channelCases.append("0.0")
                    else:
                        channelCases.append(simplify_fraction(cases, total))

                    stateChannelMatrix[i + 1].append(channelCases)

            return stateChannelMatrix
        else:
            raise ValueError("La matriz está vacía.")
    except Exception as e:
        root = tk.Tk()
        root.withdraw()  # Para evitar que se muestre la ventana Tk vacía
        messagebox.showerror("Error", str(e))
        root.destroy()
# Función para crear la matriz estado estado
def stateState(matrix):
    """
    Calculates the state-state matrix based on the given matrix.

    Args:
        matrix (list): The input matrix.

    Returns:
        list: The state-state matrix.
    """
    # Comprobar si la matriz está vacía
    if not matrix:
        messagebox.showerror("Error", "La matriz está vacía")
        return

    invertedMatrix = invertMatrix(matrix)
    invertedStateMatrix = invertMatrix(createStateMatrix(matrix))

    root = tk.Tk()
    root.withdraw()  # Para evitar que se muestre la ventana Tk vacía

    t = simpledialog.askinteger("Input", "En qué estado siguiente t+n o previo t-n desea comparar los canales? (Ponga negativo para previos)")

    root.destroy()  # Destruir la ventana Tk después de obtener la entrada



    stateStateMatrix = createStMatrix(invertedStateMatrix)

    for h in range(len(invertedStateMatrix)):
        for i in range(len(invertedStateMatrix)):
            stateCases = []
            total, cases = 0, 0
            for j in range(len(invertedMatrix)-1):
                if invertedStateMatrix[i] == invertedMatrix[j]:
                    total += 1
                    if j + t >= len(invertedMatrix):
                        if invertedMatrix[t-1] == invertedStateMatrix[h]:
                            cases += 1
                    elif invertedMatrix[j+t] == invertedStateMatrix[h]:
                        cases += 1

            # Evitar la división por 0
            if total == 0:
                stateCases.append("0.0")
            else:
                stateCases.append(simplify_fraction(cases, total))

            stateStateMatrix[i+1].append(stateCases)

    
    result_list = stateStateMatrix[1:]

    matrizestado=[]

    matrizestado.append(generate_binary_combinations(len(channels)))

    for i in range(len(result_list)):
        matrizestado.append(result_list[i])

    # Retornar la lista que no contiene 'Estados: ...'
    
    return matrizestado
# Función que genera combinaciones binarias
def generate_binary_combinations(n):
    """
    Generate all possible binary combinations of length n.

    Args:
        n (int): The length of the binary combinations.

    Returns:
        list: A list of lists, where each inner list represents a binary combination.
    """
    return [list(p) for p in product([0, 1], repeat=n)]
# Función para crear la matriz estado estado
def createStMatrix(matrix):
    """
    Creates a state matrix based on the given matrix.

    Args:
        matrix (list): The input matrix.

    Returns:
        list: The created state matrix.
    """

    createdMatrix = []
    phrase = "Estados: " + ("   ")

    for i in range(len(matrix)):
        phrase = phrase + str(matrix[i]) + " "

    createdMatrix.append(phrase)

    for i in range(len(matrix)):
        createdMatrix.append([matrix[i], ])
    return createdMatrix
# Función para dividir una lista en tres listas
def splitList(matrix):
    """
    Splits a matrix into its first title, second title, and content.

    Args:
        matrix (list): The matrix to be split.

    Returns:
        tuple: A tuple containing the first title, second title, and content.
    """
    firstTitle = matrix[0]
    secondTitle = matrix[0]
    content = convertContent(matrix[1:])
    return firstTitle, secondTitle, content
# Función para convertir el contenido de una lista
def convertContent(matrix):
    """
    Converts the content of a matrix by removing the first element of each inner list and rounding the remaining elements to 2 decimal places.

    Args:
        matrix (list): The matrix to be converted.

    Returns:
        list: The converted matrix.
    """
    newMatrixD = []
    for listO in matrix:
        newList = []
        del listO[0]
        for smolList in listO:
            for j in smolList:
                newList.append(round(j, 2))
        newMatrixD.append(newList)
        newList = []
    return newMatrixD
# Función para convertir los estados
def convertStates(State, Eval):
    """
    Converts the given State matrix based on the selected indexes in Eval.

    Args:
        State (list): The original State matrix.
        Eval (str): A string containing the selected indexes separated by spaces.

    Returns:
        list: The new matrix with values from the original State matrix based on the selected indexes.
    """
    newMatrix = []
    Indexes = []
    selectedIndexes = Eval
    selectedIndexes = list(map(int, selectedIndexes.split()))
    for i in State:
        for j in selectedIndexes:
            Indexes.append(round(i[j - 1], 2))
        newMatrix.append(Indexes)
        Indexes = []
    return newMatrix
# Función para convertir las columnas
def convertColumns2(firstState, content):
    """
    Converts the columns of a 2D list by aggregating values based on the unique elements in the firstState list.

    Args:
        firstState (list): The list containing the unique elements.
        content (list): The 2D list containing the values to be aggregated.

    Returns:
        tuple: A tuple containing the modified firstState list and the aggregated content list.
    """
    firstStateCopy = []
    contentCopy = [[] for _ in content]
    count = 0
    while count < len(firstState):
        if firstState[count] in firstStateCopy:
            index = firstStateCopy.index(firstState[count])
            for i in range(len(content)):
                row = content[i]
                rowCopy = contentCopy[i]
                rowCopy[index] = round(rowCopy[index] + row[count], 2)
        else:
            firstStateCopy.append(firstState[count])
            for i in range(len(content)):
                row = content[i]
                rowCopy = contentCopy[i]
                rowCopy.append(round(row[count], 2))
        count += 1
    return firstStateCopy, contentCopy
# Función para convertir las filas
def convertRows2(secondState, content):
    """
    Converts rows of data based on the second state and content provided.

    Args:
        secondState (list): List of second states.
        content (list): List of content.

    Returns:
        tuple: A tuple containing the updated stateCopy and contentCopy.

    """
    stateCopy = []
    contentCopy = []
    count = 0
    div = 1
    while count < len(secondState):
        if secondState[count] in stateCopy:
            if secondState[count] == secondState[0]:
                div += 1
            index = stateCopy.index(secondState[count])
            row = contentCopy[index]
            row2 = content[count]
            contentCopy[index] = [round(x + y, 2) for x, y in zip(row, row2)]
        else:
            stateCopy.append(secondState[count])
            contentCopy.append(content[count])
        count += 1
    for i in range(len(contentCopy)):
        for j in range(len(contentCopy[i])):
            contentCopy[i][j] /= div
            contentCopy[i][j] = round(contentCopy[i][j], 2)
    return stateCopy, contentCopy
# Función para marginalizar
def Marginalizar_MP(matrizest,Eval1, Eval2):
    """
    Marginalizes a given matrix based on two specified evaluations.

    Parameters:
    matrizest (list): The input matrix to be marginalized.
    Eval1 (str): The first evaluation to be considered for marginalization.
    Eval2 (str): The second evaluation to be considered for marginalization.

    Returns:
    list: The marginalized matrix.
    """
    global matrix_strMr, filaAgraficarMr
    firstState, secondState, content = splitList(matrizest)
    firstState= convertStates(firstState, Eval1)
    secondState= convertStates(secondState, Eval2)
    newFirstState, contentnuevo= convertColumns2(firstState,content)
    newSecondState, finalContent = convertRows2(secondState,contentnuevo)
    #matrizMarg = []
    #matrizMarg.append([newFirstState])
    matrizMarg=([newFirstState])
    datoComparacion= simpledialog.askstring("Datos","Escriba el estado que desea consultar")
    datoComparacion_list = list(datoComparacion)
    datoComparacion_list = [int(i) for i in datoComparacion_list]
    while len(datoComparacion)!=len(newSecondState[0]) or not datoComparacion_list in newSecondState:
        messagebox.showerror('Error', 'El estado ingresado a consultar no es valido o no se encuentra definido')
        datoComparacion= simpledialog.askstring("Datos","Escriba el estado que desea consultar")
        datoComparacion_list = list(datoComparacion)
        datoComparacion_list = [int(i) for i in datoComparacion_list]

    # Convierte newSecondState[0] a una lista de Python para poder usar el método index
    newSecondState_list = list(newSecondState)

    # Encuentra el índice de datoComparacion_list en matrixoriginalstates_list
    state_index = newSecondState_list.index(datoComparacion_list)

    filaAgraficarMr = finalContent[state_index]

    # Convertir cada sublista en una cadena y agregarla a 'matrix_str'
    matrix_strMr = [''.join(map(str, sublist)) for sublist in newFirstState]
    
        
    for i in range(len(newSecondState)):
        sublista = [newSecondState[i]]
        for number in finalContent[i]:
            sublista.append([number])
        matrizMarg.append(sublista)

    return(matrizMarg)
# Función para graficar
def graficar():
    """
    Function to plot a bar graph based on the global variables 'matrix_strMr' and 'filaAgraficarMr'.
    If these variables are not defined, an error message is displayed.
    """
    # Comprobar si las variables están definidas
    if 'matrix_strMr' not in globals() or 'filaAgraficarMr' not in globals():
        messagebox.showerror("Error", "Sin gráfica disponible, primero marginalizar.")
        return

    # Crea la figura y los ejes
    fig, ax = plt.subplots()

    # Dibuja el gráfico de barras
    ax.bar(matrix_strMr, filaAgraficarMr, width=0.5, align='center')

    # Cambia las etiquetas del eje y a [0, 0.5, 1]
    ax.set_yticks([0, 0.5, 1])

    # Guarda los recuentos del histograma en la variable resultado_original
    resultado_originalMr = filaAgraficarMr

    # Muestra el gráfico
    plt.show()
# Función para calcular las combinaciones de los estados
def Combinations(Eval1,Eval2):
    """
    Generate combinations of elements from two lists.

    Args:
        Eval1 (list): The first list of elements.
        Eval2 (list): The second list of elements.

    Returns:
        tuple: A tuple containing the combination matrix and the selected indexes.
    """
    future= Eval1
    actual= Eval2
    future = list(map(int, future))
    actual= list(map(int, actual))
    selectedIndexes=[]
    secondListCombinations = [list(comb) for longitud in range(1, len(actual) + 1) for comb in combinations(actual, longitud)]
    firstListCombinations = [list(comb) for longitud in range(1, len(future) + 1) for comb in combinations(future, longitud)]
    secondListCombinations.insert(0,[0])
    firstListCombinations.insert(0,[0])
    permutations = list(product(firstListCombinations, secondListCombinations))
    half1 = permutations[:len(permutations)//2]
    half2 = permutations[len(permutations)//2:]
    matrix=[]
    half2.reverse()
    for i in range(0, len(half1)):
        matrix.append([half1[i], half2[i]])
    listA =[]
    listB=[]
    listC=[]
    combMatrix = []
    for row in matrix:
        for lista in row:
            for tupla in lista:
                for index in tupla:
                    listA.append(index)
                listC.append(listA)
                listA=[]
            listB.append(listC)
            listC=[]
        combMatrix.append(listB)
        listB=[]
    selectedIndexes.append(future)
    selectedIndexes.append(actual)
    return combMatrix, selectedIndexes
# Función para convertir los estados de las matrices
def convertStates2(matrixState, listaindexs):
    """
    Converts the matrix state by performing the following steps:
    1. Creates a deep copy of the matrix state.
    2. Splits the copied matrix state into firstTitle, secondTitle, and content.
    3. Converts the firstTitle and secondTitle using the provided listaindexs.
    4. Converts the columns of the matrix by rearranging them based on the converted firstTitle and secondTitle.
    
    Args:
        matrixState (list): The matrix state to be converted.
        listaindexs (list): A list containing the indices for converting firstTitle and secondTitle.
    
    Returns:
        tuple: A tuple containing the final converted firstTitle, secondTitle, and content.
    """
    copiamatrixState = copy.deepcopy(matrixState)
    firstTitle, secondTitle, content = splitList(copiamatrixState)
    firstTitle = convertTitles(firstTitle, listaindexs[0])
    secondTitle = convertTitles(secondTitle, listaindexs[1])
    finalFirstTitle, finalSecondTitle, finalContent = convertColumns(firstTitle, content, secondTitle)
    return finalFirstTitle, finalSecondTitle, finalContent
# Función para convertir los títulos de las matrices
def convertTitles(titulo, selectedIndexes):
    """
    Converts the titles of a matrix based on the selected indexes.

    Args:
        titulo (list): The original matrix of titles.
        selectedIndexes (list): The list of selected indexes.

    Returns:
        list: The new matrix of titles with selected indexes.

    Example:
        >>> titulo = [['A', 'B', 'C'], ['D', 'E', 'F'], ['G', 'H', 'I']]
        >>> selectedIndexes = [1, 3]
        >>> convertTitles(titulo, selectedIndexes)
        [['A', 'C'], ['D', 'F'], ['G', 'I']]
    """
    newMatrixo = []
    indexs = []
    if selectedIndexes[0] == 0:
        return selectedIndexes
    if len(titulo[0]) == len(selectedIndexes):
        return titulo
    for i in titulo:
        for j in selectedIndexes:
            indexs.append(i[j-1])
        newMatrixo.append(indexs)
        indexs = []
    return newMatrixo
# Función para convertir las columnas de las matrices    
def convertColumns(firstTitle,content, secondTitle):
    """
    Converts the columns of a table by rearranging the data based on the given firstTitle and secondTitle.

    Args:
        firstTitle (list): The list of first titles representing the original column order.
        content (list): The list of content representing the table data.
        secondTitle (list): The list of second titles representing the desired column order.

    Returns:
        list: The converted table data with columns rearranged based on the secondTitle.
    """
    
    firstTitleCopy = []
    contentCopy = []
    rowcopia = []
    count = 0
    if firstTitle[0]==0:
        return sumColumns(firstTitle,content, secondTitle)
    else:
        for i in content:
                contentCopy.append([])
        while(count < len(firstTitle)):
            if(firstTitle[count] in firstTitleCopy ):
                index = firstTitleCopy.index(firstTitle[count])
                for i in range(len(content)):
                    row= content[i]
                    rowcopia = contentCopy[i]
                    rowcopia[index] = rowcopia[index]+row[count]
            else:
                firstTitleCopy.append(firstTitle[count])
                for i in range (len(content)):
                    row = content[i]
                    rowcopia = contentCopy[i]
                    rowcopia.append(row[count])
            count +=1
        return convertRows(secondTitle, contentCopy, firstTitleCopy)
# Función para convertir las filas de las matrices
def convertRows(secondTitle, content, firstTitlefinal):
    """
    Converts rows in the content based on the secondTitle and returns the modified content.

    Parameters:
    secondTitle (list): List of second titles.
    content (list): List of content rows.
    firstTitlefinal (str): First title.

    Returns:
    tuple: A tuple containing the firstTitlefinal, modified titleCopy, and modified contentCopy.
    """
    titleCopy = []
    contentCopy = []
    count = 0
    div = 1
    if secondTitle[0] == 0:
        return sumRows(secondTitle, content, firstTitlefinal)
    else:
        while count < len(secondTitle):
            if secondTitle[count] in titleCopy:
                if secondTitle[count] == secondTitle[0]:
                    div += 1
                index = titleCopy.index(secondTitle[count])
                row = contentCopy[index]
                secondRow = content[count]
                contentCopy[index] = [(x + y) for x, y in zip(row, secondRow)]
            else:
                titleCopy.append(secondTitle[count])
                contentCopy.append(content[count])
            count += 1
        for i in range(len(contentCopy)):
            for j in range(len(contentCopy[i])):
                contentCopy[i][j] /= div
                contentCopy[i][j] = round(contentCopy[i][j], 2)
        return firstTitlefinal, titleCopy, contentCopy
# Función para calcular la suma de las columnas de una matriz
def sumColumns(firstTitle,content,secondTitle):
    """
    Sums the values in each column of the given content and returns the result as a list of rows.

    Args:
        firstTitle (str): The title of the first column.
        content (list): The content containing the values in each column.
        secondTitle (str): The title of the second column.

    Returns:
        list: The result of summing the values in each column, represented as a list of rows.
    """
    suma_rows = []
    for row in content:
        suma_row = sum(row)
        suma_rows.insert(0,[suma_row]) 
    suma_rows.reverse()
    return convertRows(secondTitle,suma_rows, firstTitle)
# Función para calcular la suma de las filas de una matriz
def sumRows(secondTitle,content, firstTitlefinal):
    """
    Calculate the sum of each column in the content list and return the result.

    Parameters:
    secondTitle (str): The second title.
    content (list): A list of lists representing the content.
    firstTitlefinal (str): The first title.

    Returns:
    tuple: A tuple containing the first title, second title, and the sum of each column in the content list.
    """
    div=len(content)
    content_total = [sum(column) for column in zip(*content)]
    content_total = [round(elemento / div,2) for elemento in content_total]
    return (firstTitlefinal,secondTitle, content_total)
# Función para calcular el producto de Kronecker de una lista de matrices
def kronecker_product_matrices(matrix_list):
    """
    Computes the Kronecker product of a list of matrices.

    Args:
        matrix_list (list): A list of matrices.

    Returns:
        matrix: The Kronecker product of the matrices in the list.
    """
    result = matrix_list[0]
    for matrix in matrix_list[1:]:
        result = kronecker_product(result, matrix)
    return result
# Función para calcular el producto de Kronecker de dos matrices
def kronecker_product(A, B):
    """
    Computes the Kronecker product of two matrices A and B.

    Parameters:
    A (numpy.ndarray): The first matrix of shape (m, n).
    B (numpy.ndarray): The second matrix of shape (p, q).

    Returns:
    numpy.ndarray: The Kronecker product of A and B, a matrix of shape (m * p, n * q).
    """
    m, n = A.shape
    p, q = B.shape

    result = [[0] * (n * q) for _ in range(m * p)]

    for i in range(m):
        for j in range(n):
            for k in range(p):
                for l in range(q):
                    result[i * p + k][j * q + l] = A[i, j] * B[k, l]

    return np.array(result)
# Función para calcular la distancia EMD entre dos histogramas
def earth_movers_distance(hist1, hist2):
    """
    Calculates the Earth Mover's Distance between two histograms.

    Parameters:
    hist1 (numpy.ndarray): The first histogram.
    hist2 (numpy.ndarray): The second histogram.

    Returns:
    float: The Earth Mover's Distance between the two histograms.
    """
    # Rellena el histograma más corto con ceros hasta que tenga la misma longitud que el histograma más largo
    if len(hist1) < len(hist2):
        hist1 = np.pad(hist1, (0, len(hist2) - len(hist1)), 'constant')
    elif len(hist2) < len(hist1):
        hist2 = np.pad(hist2, (0, len(hist1) - len(hist2)), 'constant')

    n = len(hist1)
    total_distance = 0.0
    cumulative_diff = 0.0

    for i in range(n):
        diff = hist1[i] - hist2[i]
        cumulative_diff += diff
        total_distance += abs(cumulative_diff)

    # Redondear el resultado a 2 decimales
    return round(total_distance, 2)
# Función para calcular la distancia EMD entre dos histogramas
def DisEMD(matrixState,Eval1,Eval2):
    """
    Calculates the Earth Mover's Distance (EMD) between two sets of evaluations.

    Args:
        matrixState (list): The matrix state.
        Eval1 (list): The first set of evaluations.
        Eval2 (list): The second set of evaluations.

    Returns:
        str: A string containing the minimum EMD and the corresponding combination.
    """

    global matrix_strMD, filaAgraficarMD, min_histogramMD, filaAgraficarOMD


    combMatrix, primer_matrix = Combinations(Eval1,Eval2)
    matrixoriginal=convertStates2(matrixState,primer_matrix)
    matrixoriginalstates=matrixoriginal[1:2]
    matrixoriginalfirsttitle=matrixoriginal[0:1]
    matrixoriginal=np.array(matrixoriginal[2:])
    datoComparacion= simpledialog.askstring("Datos","Escriba el estado que desea consultar")
    datoComparacion_list = list(datoComparacion)
    datoComparacion_list = [int(i) for i in datoComparacion_list]
    while len(datoComparacion)!=len(matrixoriginalstates[0][0]) or not datoComparacion_list in matrixoriginalstates[0]:
        messagebox.showerror('Error', 'El estado ingresado a consultar no es valido o no se encuentra definido')
        datoComparacion= simpledialog.askstring("Datos","Escriba el estado que desea consultar")
        datoComparacion_list = list(datoComparacion)
        datoComparacion_list = [int(i) for i in datoComparacion_list]

    # Convierte matrixoriginalstates[0] a una lista de Python para poder usar el método index
    matrixoriginalstates_list = list(matrixoriginalstates[0])

    # Encuentra el índice de datoComparacion_list en matrixoriginalstates_list
    state_index = matrixoriginalstates_list.index(datoComparacion_list)

    filaAgraficarOMD = matrixoriginal[0][state_index]

    # Convertir cada sublista en una cadena y agregarla a 'matrix_str'
    matrix_strMD = [''.join(map(str, sublist)) for sublist in matrixoriginalfirsttitle[0]]

    # Guarda los recuentos del histograma en la variable resultado_original
    resultado_original = filaAgraficarOMD


    counti,countj=0,0

    # Inicializa el arreglo para almacenar los resultados del histograma
    histogram_results = []

    while (counti<len(combMatrix)):
        matrix1=combMatrix[counti][countj]
        matrix2=combMatrix[counti][countj+1]
        matrix1=convertStates2(matrixState,matrix1)[2:]
        matrix2=convertStates2(matrixState,matrix2)[2:]
        
        if not isinstance(matrix1, list):
            matrix1 = [matrix1]
        if not isinstance(matrix2, list):
            matrix2 = [matrix2]

        A=np.array(matrix1)
        B=np.array(matrix2)
        
        if A.ndim != 2:
            A = A.reshape((1, -1))
        if B.ndim != 2:
            B = B.reshape((1, -1))
        
        result = kronecker_product_matrices([A, B])
        # Convertir 'result' en una lista de listas
        result_list = result.tolist()
        # Dividir cada sublista en sublistas de longitud 'len(matrix_str)' y redondear cada elemento a 2 decimales
        new_result_list = [[round(num, 2) for num in sublist[i:i+len(matrix_strMD)]] for sublist in result_list for i in range(0, len(sublist), len(matrix_strMD))]

        filaAgraficarMD = new_result_list[state_index]

        # Guarda los conteos en un arreglo
        histogram_results.append(filaAgraficarMD)
        
        counti+=1

    
    EMDresults = []
    for i in range(1, len(histogram_results)):  # Comienza desde el segundo elemento
        EMDresults.append(earth_movers_distance(resultado_original, histogram_results[i]))
    combinaciones = []
    for combinacion in combMatrix[1:]:  # Comienza desde el segundo elemento
        combinaciones.append(combinacion[0] + combinacion[1])

    # Encuentra el índice de la menor distancia EMD
    min_index = np.argmin(EMDresults)

    # Encuentra la menor distancia EMD y la combinación correspondiente
    min_emd = EMDresults[min_index]
    min_combinacion = combinaciones[min_index]


    # Encuentra el histograma correspondiente
    min_histogramMD = histogram_results[min_index + 1]  # Ajusta el índice para tener en cuenta que comenzamos desde el segundo elemento


    return (f"Menor distancia EMD: {min_emd}, Combinación correspondiente: {min_combinacion}")
# Función para graficar el histograma original
def graficarO():
    """
    Function to plot a bar graph.

    This function checks if the required variables are defined and displays an error message if not.
    It creates a figure and axes, and then plots a bar graph using the given data.
    The y-axis labels are set to [0, 0.5, 1].
    Finally, it displays the graph.
    """
    # Comprobar si las variables están definidas
    if 'matrix_strMD' not in globals() or 'filaAgraficarOMD' not in globals():
        messagebox.showerror("Error", "Sin gráfica disponible, hallar la distancia EMD.")
        return

    # Crea la figura y los ejes
    fig, ax = plt.subplots()

    # Dibuja el gráfico de barras
    ax.bar(matrix_strMD, filaAgraficarOMD, width=0.5, align='center')

    # Cambia las etiquetas del eje y a [0, 0.5, 1]
    ax.set_yticks([0, 0.5, 1])

    # Muestra el gráfico
    plt.show()
# Función para graficar el histograma ganador
def graficarG():
    """
    Plot the histogram corresponding to the shortest EMD distance.

    Checks if the variables 'matrix_strMD', 'rowAgraficarMD' and 'min_histogramMD' are defined.
    If any of these variables are not defined, it displays an error message and returns.

    Create a shape and axes for the graph.
    Draws the bar chart using the values of 'matrix_strMD' and 'min_histogramMD'.
    Change the y-axis labels to [0, 0.5, 1].
    Set the graph title to 'Histogram corresponding to the shortest EMD distance'.
    Shows the graph.
    """
    if 'matrix_strMD' not in globals() or 'filaAgraficarMD' not in globals() or 'min_histogramMD' not in globals():
        messagebox.showerror("Error", "Sin gráfica disponible, hallar la distancia EMD.")
        return

    fig, ax = plt.subplots()

    ax.bar(matrix_strMD, min_histogramMD, width=0.5, align='center')

    ax.set_yticks([0, 0.5, 1])

    plt.title('Histograma correspondiente a la menor distancia EMD')

    plt.show()
# Funciones para hacer el llamado de cada función, asi que la documentación es la misma
def Ingresar_B():

    def Crear_B():
        global channels
        filas_str = cuadro_canales.get()
        columnas_str = cuadro_muestras.get()
        try:
            filas = int(filas_str)
            columnas = int(columnas_str)
            if filas > 0 and columnas > 0:  # Verificar si los valores son mayores que 0
                limpiar_frameTr()  # Limpiar el frame solo si los valores son válidos
                crear_tabla(filas, columnas)
            else:
                messagebox.showerror("Error", "Los campos de canales y muestras deben ser mayores que 0.")
        except ValueError:
            messagebox.showerror("Error", "Los campos de canales y muestras deben ser números enteros.")


    limpiar_frameTr()
    CantCanales = tk.Label(frame_trabajo, text="Cantidad de Canales", bg="#434343", fg="white", font=('Helvetica', 15))
    CantCanales.place(relx=0.4, rely=0.4, anchor="center")

    cuadro_canales = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    cuadro_canales.place(relx=0.65, rely=0.4, anchor="center")

    CantMuestras = tk.Label(frame_trabajo, text="Cantidad de Muestras", bg="#434343", fg="white", font=('Helvetica', 15))
    CantMuestras.place(relx=0.4, rely=0.5, anchor="center")

    cuadro_muestras = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    cuadro_muestras.place(relx=0.65, rely=0.5, anchor="center")

    crearM_Button=tk.Button(frame_trabajo, text="Crear Matriz", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15), command=Crear_B)
    crearM_Button.place(relx=0.65,rely=0.6, anchor="center")

def Cargar_B():

    def Ruta_Archivo():
        global channels
        file_path = cuadro_ruta.get()
        if file_path:
            if os.path.isfile(file_path):
                channels = loadFromFile(file_path)
                if channels:
                    volver_a_principal()
            else:
                messagebox.showerror("Error", "La ruta del archivo no es correcta.")
        else:
            messagebox.showerror("Error", "La ruta del archivo no puede estar vacía.")

    def Cargar_Archivo():
        global channels
        channels = abrir_explorador_y_cargar_archivo()
        if channels:
            volver_a_principal()

    limpiar_frameTr()
    Ruta = tk.Label(frame_trabajo, text="Ingresar ruta", bg="#434343", fg="white", font=('Helvetica', 15))
    Ruta.place(relx=0.4, rely=0.5, anchor="center")

    cuadro_ruta = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    cuadro_ruta.place(relx=0.6, rely=0.5, anchor="center")

    ruta_Button=tk.Button(frame_trabajo,text="Subir Archivo", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=Ruta_Archivo)
    ruta_Button.place(relx=0.6,rely=0.6, anchor="center")   

    Subir_Button=tk.Button(frame_trabajo,image=icono_tk_opcion1, bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=Cargar_Archivo)
    Subir_Button.place(relx=0.75,rely=0.5, anchor="center")         

def MatrizO_B():
    limpiar_frameTr()
    imprimir_tabla(channels)

def MatrizC_B():
    limpiar_frameTr()
    a=channelState(channels)
    imprimir_tablaCS(a)

def MatrizE_B():
    limpiar_frameTr()
    global matrixOrE
    matrixOrE = stateState(channels)

    # Insertar "Estados" en la posición [0][0] y desplazar los demás elementos a la derecha
    if matrixOrE:
        matrixOrE[0].insert(0, "Estados")

        imprimir_tablaS(matrixOrE)
        matrixOrE[0].pop(0)

def IngMatrizE_B():

    limpiar_frameTr()

    def Ruta_ArchivoMatE():
        global matrixOrE
        file_path = cuadro_ruta.get()
        if file_path:
            if os.path.isfile(file_path):
                matrixOrE = loadFromFileMatrixE(file_path)
                if matrixOrE:
                    matrixOrE[0].insert(0, "Estados")

                    imprimir_tablaS(matrixOrE)
                    matrixOrE[0].pop(0)                    
                    
            else:
                messagebox.showerror("Error", "La ruta del archivo no es correcta.")
        else:
            messagebox.showerror("Error", "La ruta del archivo no puede estar vacía.")

    def Cargar_ArchivoMatE():
        global matrixOrE
        matrixOrE = abrir_explorador_y_cargar_archivoMatrixE()
        
        if matrixOrE:
            matrixOrE[0].insert(0, "Estados")

            imprimir_tablaS(matrixOrE)
            matrixOrE[0].pop(0)


    Ruta = tk.Label(frame_trabajo, text="Ingresar ruta", bg="#434343", fg="white", font=('Helvetica', 15))
    Ruta.place(relx=0.4, rely=0.5, anchor="center")

    cuadro_ruta = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    cuadro_ruta.place(relx=0.6, rely=0.5, anchor="center")

    ruta_Button=tk.Button(frame_trabajo,text="Subir Archivo", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=Ruta_ArchivoMatE)
    ruta_Button.place(relx=0.6,rely=0.6, anchor="center")   

    Subir_Button=tk.Button(frame_trabajo,image=icono_tk_opcion1, bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=Cargar_ArchivoMatE)
    Subir_Button.place(relx=0.75,rely=0.5, anchor="center")       

def Marginalizar_B():
    limpiar_frameTr()


    def MargiP_B():
        try:
            # Check if matrixSB exists
            matrixSB = copy.deepcopy(matrixOrE)
        except NameError:
            messagebox.showerror("Error", "Primero debe crear la matriz estado.")
            return

        # Get the values
        cp_estados = CP_Estados.get()
        cs_estados = CS_Estados.get()

        try:
            # Check if the values are empty
            if not cp_estados or not cs_estados:
                raise ValueError("Los valores no pueden estar vacíos.")

            # Check if the values are positive integers and separated by a space
            if not all(i.isdigit() and int(i) >= 0 for i in cp_estados.split()) or not all(i.isdigit() and int(i) >= 0 for i in cs_estados.split()):
                raise ValueError("Los valores ingresados deben ser enteros positivos y deben estar separados por un espacio.")

            # Check if the values are not a single string of numbers
            if len(cp_estados) > 1 and " " not in cp_estados or len(cs_estados) > 1 and " " not in cs_estados:
                raise ValueError("Los valores ingresados deben estar separados por un espacio.")

            # Check if the length of the characters entered does not exceed len(matrixSB[0][0]) plus the respective empty spaces for each entry
            if len(cp_estados) > len(matrixSB[0][0]) + cp_estados.count(' ') or len(cs_estados) > len(matrixSB[0][0]) + cs_estados.count(' '):
                raise ValueError("La longitud de los caracteres ingresados no debe exceder la longitud de los canales de la matriz.")

            a = Marginalizar_MP(matrixSB, cp_estados, cs_estados)
            a[0].insert(0, "Estados")
            imprimir_tablaS(a)

        except ValueError as e:
            messagebox.showerror("Error", e)
 


    P_Estados = tk.Label(frame_trabajo, text="Estado futuro", bg="#434343", fg="white", font=('Helvetica', 15))
    P_Estados.place(relx=0.4, rely=0.4, anchor="center")

    CP_Estados = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    CP_Estados.place(relx=0.65, rely=0.4, anchor="center")

    S_Estados = tk.Label(frame_trabajo, text="Estado presente", bg="#434343", fg="white", font=('Helvetica', 15))
    S_Estados.place(relx=0.4, rely=0.5, anchor="center")

    CS_Estados = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    CS_Estados.place(relx=0.65, rely=0.5, anchor="center")

    Margi_Button=tk.Button(frame_trabajo,text="Marginzalizar", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=MargiP_B)
    Margi_Button.place(relx=0.65,rely=0.6, anchor="center")       

def Grafica_B():
    limpiar_frameTr()
    graficar()

def DistanciaE_B():
    limpiar_frameTr()

    def Dis_B():

        try:
            # Check if matrixSB exists
            matrixSB = copy.deepcopy(matrixOrE)
        except NameError:
            messagebox.showerror("Error", "Primero debe crear la matriz estado.")
            return

        # Get the values
        cp_estados = CP_Estados.get()
        c_dato = C_Dato.get()

        try:
            # Check if the values are empty
            if not cp_estados or not c_dato:
                raise ValueError("Los campos no pueden estar vacíos.")

            # Check if the values are positive integers
            if not all(i.isdigit() and int(i) >= 0 for i in cp_estados.split()) or not all(i.isdigit() and int(i) >= 0 for i in c_dato.split()):
                raise ValueError("Los valores ingresados deben ser enteros positivos.")

            # Check if the length of the characters entered does not exceed len(matrixSB[0][0])
            if len(cp_estados) > len(matrixSB[0][0]) or len(c_dato) > len(matrixSB[0][0]):
                raise ValueError("La longitud de los caracteres ingresados no debe exceder la longitud de los canales de la matriz.")

            limpiar_frameTr()
            printo = DisEMD(matrixSB, cp_estados, c_dato)
            L_Printo = tk.Label(frame_trabajo, text=printo, bg="#434343", fg="white", font=('Helvetica', 18))
            L_Printo.place(relx=0.5, rely=0.5, anchor="center")

        except ValueError as e:
            messagebox.showerror("Error", e)


    P_Estados = tk.Label(frame_trabajo, text="Estado Futuro", bg="#434343", fg="white", font=('Helvetica', 15))
    P_Estados.place(relx=0.4, rely=0.4, anchor="center")

    CP_Estados = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    CP_Estados.place(relx=0.65, rely=0.4, anchor="center")

    S_Estados = tk.Label(frame_trabajo, text="Estado Presente", bg="#434343", fg="white", font=('Helvetica', 15))
    S_Estados.place(relx=0.4, rely=0.5, anchor="center")

    C_Dato = tk.Entry(frame_trabajo, bg="#332F2C", fg="white", font=('Helvetica', 15), bd=2, relief=tk.GROOVE)
    C_Dato.place(relx=0.65, rely=0.5, anchor="center")    

    Margi_Button=tk.Button(frame_trabajo,text="Calcular EMD", bg="#332F2C", fg="white", padx=5, pady=5, font=('Helvetica', 15),command=Dis_B)
    Margi_Button.place(relx=0.65,rely=0.6, anchor="center")              

def GraficaO_B():
    limpiar_frameTr()
    graficarO()

def GraficaG_B():
    limpiar_frameTr()
    graficarG()

# Diccionaro para la conversión de letras a números
diccionario_letras = {}
for numero in range(1, 27):
    letra = chr(ord('A') + numero - 1)
    diccionario_letras[numero] = letra


# Configuración de la ventana principal
ventana_principal = tk.Tk()
ventana_principal.title("Análisis de probabilidades")
ventana_principal.geometry("1200x720")
ventana_principal.configure(bg="#332F2C")

# Configuración de la barra lateral
barra_lateral = tk.Frame(ventana_principal, bg="#292523", bd=2)
barra_lateral.pack(side=tk.LEFT, fill=tk.Y)

contenido_label = tk.Label(ventana_principal, text="Análisis de probabilidades", bg="#332F2C", fg="white", padx=20, pady=20, font=('Helvetica', 25))
contenido_label.pack()

# Configuración del frame de trabajo

barra_superior = tk.Frame(ventana_principal, bg="#332F2C", bd=5, relief=tk.GROOVE, height=45)
barra_superior.pack(padx=10, side=tk.TOP, fill=tk.X)

About_Button = tk.Button(barra_superior, text="About", bg="#332F2C", fg="white", padx=45, pady=5, font=('Helvetica', 10),command=aboutB)
About_Button.pack(side=tk.LEFT)

frame_trabajo = tk.Frame(ventana_principal, bg="#434343", bd=5, relief=tk.GROOVE)
frame_trabajo.pack(padx=10,pady=10, fill=tk.BOTH, expand=True)



# Configuración de la imagen en la barra lateral
ruta_imagen_superior = r"Pic.jpg"
imagen_pil_superior = Image.open(ruta_imagen_superior)
imagen_pil_superior = imagen_pil_superior.resize((200, 200))
imagen_tk_superior = ImageTk.PhotoImage(imagen_pil_superior)

imagen_label_superior = tk.Label(barra_lateral, image=imagen_tk_superior, bg="#292523")
imagen_label_superior.image = imagen_tk_superior
imagen_label_superior.pack(pady=20)

# Configuración de los íconos para los botones
icono_principal = Image.open(r"Íconos\Principal.png").resize((30, 30))
icono_opcion1 = Image.open(r"Íconos\Ingresar.png").resize((30, 30))
icono_opcion2 = Image.open(r"Íconos\Analisis.png").resize((30, 30))
icono_opcion3 = Image.open(r"Íconos\Visualizar.png").resize((30, 30))
icono_opcion4 = Image.open(r"Íconos\Fallo.png").resize((30, 30))
icono_salir = Image.open(r"Íconos\Salir.png").resize((30, 30))

icono_tk_principal = ImageTk.PhotoImage(icono_principal)
icono_tk_opcion1 = ImageTk.PhotoImage(icono_opcion1)
icono_tk_opcion2 = ImageTk.PhotoImage(icono_opcion2)
icono_tk_opcion3 = ImageTk.PhotoImage(icono_opcion3)
icono_tk_opcion4 = ImageTk.PhotoImage(icono_opcion4)
icono_tk_salir = ImageTk.PhotoImage(icono_salir)

# Estilo para los botones con íconos
boton_estilo_icono = {'bg': "#292523", 'bd': '2', 'highlightthickness': 2, 'fg': 'white'}

# Configuración de los botones con íconos y texto
boton_principal = tk.Button(barra_lateral, image=icono_tk_principal, text="Principal", compound=tk.LEFT, command=volver_a_principal, anchor="w", padx=10, **boton_estilo_icono)
boton_opcion1 = tk.Button(barra_lateral, image=icono_tk_opcion1, text="Ingreso de matrices", compound=tk.LEFT, command=opcion1, anchor="w", padx=10, **boton_estilo_icono)
boton_opcion2 = tk.Button(barra_lateral, image=icono_tk_opcion2, text="Análisis de matrices asociadas", compound=tk.LEFT, command=opcion2, anchor="w", padx=10, **boton_estilo_icono)
boton_opcion3 = tk.Button(barra_lateral, image=icono_tk_opcion3, text="Visualización y probabilidades", compound=tk.LEFT, command=opcion3, anchor="w", padx=10, **boton_estilo_icono)
boton_opcion4 = tk.Button(barra_lateral, image=icono_tk_opcion4, text="Mínima ganancia", compound=tk.LEFT, command=opcion4, anchor="w", padx=10, **boton_estilo_icono)

# Nuevo botón para salir del programa
boton_salir = tk.Button(barra_lateral, image=icono_tk_salir, text="Salir del programa", compound=tk.LEFT, command=salir_del_programa, anchor="w", padx=10, **boton_estilo_icono)

# Empaque de botones en la barra lateral
boton_principal.pack(fill=tk.X)
boton_opcion1.pack(fill=tk.X)
boton_opcion2.pack(fill=tk.X)
boton_opcion3.pack(fill=tk.X)
boton_opcion4.pack(fill=tk.X)
boton_salir.pack(fill=tk.X)

# Configuración de la segunda imagen en la barra lateral (debajo de los botones)
ruta_imagen_inferior = r"Pat.jpg"
imagen_pil_inferior = Image.open(ruta_imagen_inferior)
imagen_pil_inferior = imagen_pil_inferior.resize((200, 200))
imagen_tk_inferior = ImageTk.PhotoImage(imagen_pil_inferior)

imagen_label_inferior = tk.Label(barra_lateral, image=imagen_tk_inferior, bg="#292523")
imagen_label_inferior.image = imagen_tk_inferior
imagen_label_inferior.pack(pady=20)

# Lanzar la interfaz gráfica
ventana_principal.mainloop()






